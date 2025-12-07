import logging
import json
import os
import uuid
import shutil
from typing import Dict, Any, Optional, List
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from io import BytesIO
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import func
from ..database import get_db
from .. import models
from ...ai_agent.agent import AIAgent
from ..utils.repo_context import clone_repo_to_temp as clone_repo, cleanup_repo
from ..config import settings # Keep settings as it's used later

logger = logging.getLogger(__name__)

def _is_error_response(response: str) -> bool:
    """Check if an AI response is an error message."""
    if not response:
        return True
    
    error_indicators = [
        "Error:",
        "Failed:",
        "AI analysis failed:",
        "API key",
        "rate limit",
        "connection error",
        "timeout",
        "HTTPException",
        "Authentication failed"
    ]
    
    response_lower = response.lower()
    for indicator in error_indicators:
        if indicator.lower() in response_lower:
            return True
    
    return False

router = APIRouter(
    prefix="/ai",
    tags=["ai"]
)

@router.get("/config")
async def get_ai_config():
    """Get current AI configuration."""
    return {
        "provider": settings.AI_PROVIDER,
        "model": settings.AI_MODEL,
        "openai_model": settings.OPENAI_MODEL,
        "anthropic_model": settings.ANTHROPIC_MODEL,
        "docker_model": getattr(settings, 'DOCKER_MODEL', 'ai/llama3.2:latest'),
        "docker_base_url": settings.DOCKER_BASE_URL
    }

# Initialize AI Agent
# In a real app, this might be a dependency to allow for mocking/lazy loading
try:
    # Determine model and url based on provider
    provider = settings.AI_PROVIDER
    model = settings.AI_MODEL
    ollama_url = settings.OLLAMA_BASE_URL
    
    if provider == "openai" and settings.OPENAI_MODEL:
        model = settings.OPENAI_MODEL
    elif provider in ["claude", "anthropic_foundry"] and settings.ANTHROPIC_MODEL:
        # Both claude and anthropic_foundry use ANTHROPIC_MODEL
        model = settings.ANTHROPIC_MODEL
    elif provider == "docker":
        ollama_url = settings.DOCKER_BASE_URL
        model = settings.DOCKER_MODEL or "ai/llama3.2:latest"
        
    ai_agent = AIAgent(
        openai_api_key=settings.OPENAI_API_KEY,
        anthropic_api_key=settings.ANTHROPIC_API_KEY,
        provider=provider,
        model=model,
        ollama_base_url=ollama_url,
        azure_foundry_endpoint=settings.AZURE_AI_FOUNDRY_ENDPOINT,
        azure_foundry_api_key=settings.AZURE_AI_FOUNDRY_API_KEY
    )
except Exception as e:
    logger.warning(f"Failed to initialize AI Agent: {e}")
    ai_agent = None

# Initialize Diagrams Index
from ..utils.diagrams_indexer import get_diagrams_index
diagrams_index = {}
try:
    diagrams_index = get_diagrams_index()
except Exception as e:
    logger.warning(f"Failed to initialize diagrams index: {e}")

class RemediationRequest(BaseModel):
    vuln_type: str
    description: str
    context: str
    language: str
    finding_id: Optional[str] = None

class RemediationResponse(BaseModel):
    remediation: str
    diff: str
    remediation_id: Optional[str] = None

@router.post("/remediate", response_model=RemediationResponse)
async def generate_remediation(
    request: RemediationRequest,
    db: Session = Depends(get_db)
):
    """Generate remediation for a vulnerability."""
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
    
    try:
        result = await ai_agent.generate_remediation(
            vuln_type=request.vuln_type,
            description=request.description,
            context=request.context,
            language=request.language
        )
        
        remediation_text = result.get("remediation", "")
        diff_text = result.get("diff", "")
        remediation_id = None

        # Persist if finding_id is provided
        if request.finding_id:
            try:
                # Verify finding exists
                finding = db.query(models.Finding).filter(models.Finding.id == request.finding_id).first()
                if finding:
                    new_remediation = models.Remediation(
                        finding_id=request.finding_id,
                        remediation_text=remediation_text,
                        diff=diff_text,
                        confidence=0.85 # Placeholder/Default confidence from AI
                    )
                    db.add(new_remediation)
                    
                    # Update the "latest" cache on the finding for backward compatibility
                    finding.ai_remediation_text = remediation_text
                    finding.ai_remediation_diff = diff_text
                    
                    db.commit()
                    db.refresh(new_remediation)
                    remediation_id = str(new_remediation.id)
            except Exception as db_err:
                logger.error(f"Failed to persist remediation: {db_err}")
                # Don't fail the request if persistence fails, just log it
        
        return RemediationResponse(
            remediation=remediation_text,
            diff=diff_text,
            remediation_id=remediation_id
        )
    except Exception as e:
        logger.error(f"Error generating remediation: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class TriageRequest(BaseModel):
    title: str
    description: str
    severity: str
    scanner: str
    finding_id: Optional[str] = None

class TriageResponse(BaseModel):
    priority: str
    confidence: float
    reasoning: str
    false_positive_probability: float

@router.post("/triage", response_model=TriageResponse)
async def triage_finding(
    request: TriageRequest,
    db: Session = Depends(get_db)
):
    """Analyze and triage a finding."""
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
        
    try:
        result = await ai_agent.triage_finding(
            title=request.title,
            description=request.description,
            severity=request.severity,
            scanner=request.scanner
        )
        
        reasoning = result.get("reasoning", "Analysis failed")
        
        # Append to finding description if finding_id is provided
        if request.finding_id:
            try:
                # Try to find by finding_uuid first (as that's what the API exposes as 'id')
                # But handle potential UUID format errors if string is passed
                import uuid
                try:
                    f_uuid = uuid.UUID(request.finding_id)
                    finding = db.query(Finding).filter(Finding.finding_uuid == f_uuid).first()
                    if not finding:
                        # Fallback to primary key
                        finding = db.query(Finding).filter(Finding.id == f_uuid).first()
                except ValueError:
                    finding = None
                    logger.warning(f"Invalid UUID format for finding_id: {request.finding_id}")

                if finding:
                    # Check if reasoning is already appended to avoid duplication
                    ai_header = "### 🤖 AI Security Analysis"
                    if ai_header not in (finding.description or ""):
                        # Format the reasoning nicely
                        formatted_reasoning = (
                            f"\n\n{ai_header}\n"
                            f"**Priority:** {result.get('priority', 'Medium')}\n"
                            f"**Confidence:** {result.get('confidence', 0.0) * 100:.0f}%\n\n"
                            f"{reasoning}\n"
                        )
                        finding.description = f"{finding.description or ''}{formatted_reasoning}"
                        db.commit()
                        logger.info(f"Appended AI analysis to finding {request.finding_id}")
                else:
                    logger.warning(f"Finding not found for ID: {request.finding_id}")
            except Exception as db_err:
                logger.error(f"Failed to update finding description: {db_err}")

        return TriageResponse(
            priority=result.get("priority", "Medium"),
            confidence=result.get("confidence", 0.0),
            reasoning=reasoning,
            false_positive_probability=result.get("false_positive_probability", 0.0)
        )
    except Exception as e:
        logger.error(f"Error triaging finding: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class AnalyzeFindingRequest(BaseModel):
    finding_id: str
    prompt: Optional[str] = None

@router.post("/analyze-finding")
async def analyze_finding_endpoint(
    request: AnalyzeFindingRequest,
    db: Session = Depends(get_db)
):
    """Analyze a specific finding using AI."""
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    try:
        # Fetch finding
        # Try UUID first
        try:
            f_uuid = uuid.UUID(request.finding_id)
            finding = db.query(models.Finding).filter(models.Finding.finding_uuid == f_uuid).first()
            if not finding:
                finding = db.query(models.Finding).filter(models.Finding.id == f_uuid).first()
        except ValueError:
            # If not UUID, try ID (though ID is usually UUID in this schema)
            finding = None
        
        if not finding:
            raise HTTPException(status_code=404, detail="Finding not found")

        # Convert to dict for AI
        finding_dict = {
            "title": finding.title,
            "description": finding.description,
            "severity": finding.severity,
            "file_path": finding.file_path,
            "line_start": finding.line_start,
            "line_end": finding.line_end,
            "code_snippet": finding.code_snippet,
            "finding_type": finding.finding_type,
            "scanner": finding.scanner_name
        }

        analysis = await ai_agent.analyze_finding(
            finding=finding_dict,
            user_prompt=request.prompt
        )
        
        # Append AI Security Analysis to finding description if successful
        if analysis and not _is_error_response(analysis):
            from datetime import datetime
            timestamp = datetime.utcnow().isoformat()
            ai_header = "### 🤖 AI Security Analysis"
            
            # Check if analysis is already appended to avoid duplication
            if ai_header not in (finding.description or ""):
                provider_name = getattr(ai_agent, 'provider_name', 'AI')
                formatted_analysis = f"""

{ai_header}
**Analyzed:** {timestamp}
**Provider:** {provider_name}

{analysis}

---
"""
                finding.description = f"{finding.description or ''}{formatted_analysis}"
                db.commit()
                logger.info(f"Appended AI Security Analysis to finding {request.finding_id}")
        
        return {"analysis": analysis}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing finding: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class AnalyzeComponentRequest(BaseModel):
    package_name: str
    version: str
    package_manager: str

@router.post("/analyze-component")
async def analyze_component_endpoint(
    request: AnalyzeComponentRequest,
    db: Session = Depends(get_db)
):
    """Analyze a software component using AI, with caching."""
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    try:
        # 1. Check DB for existing analysis
        existing = db.query(models.ComponentAnalysis).filter(
            models.ComponentAnalysis.package_name == request.package_name,
            models.ComponentAnalysis.version == request.version,
            models.ComponentAnalysis.package_manager == request.package_manager
        ).first()

        if existing:
            return {
                "analysis_text": existing.analysis_text,
                "vulnerability_summary": existing.vulnerability_summary,
                "severity": existing.severity,
                "exploitability": existing.exploitability,
                "fixed_version": existing.fixed_version,
                "source": "cache"
            }

        # 2. Call AI Provider
        analysis = await ai_agent.provider.analyze_component(
            package_name=request.package_name,
            version=request.version,
            package_manager=request.package_manager
        )

        # 3. Save to DB
        new_analysis = models.ComponentAnalysis(
            package_name=request.package_name,
            version=request.version,
            package_manager=request.package_manager,
            analysis_text=analysis.get("analysis_text", ""),
            vulnerability_summary=analysis.get("vulnerability_summary", ""),
            severity=analysis.get("severity", "Unknown"),
            exploitability=analysis.get("exploitability", "Unknown"),
            fixed_version=analysis.get("fixed_version", "Unknown")
        )
        
        try:
            db.add(new_analysis)
            db.commit()
            db.refresh(new_analysis)
        except Exception as db_err:
            logger.error(f"Failed to cache component analysis: {db_err}")
            # Continue even if caching fails
        
        return {
            **analysis,
            "source": "ai"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing component: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/remediate/{remediation_id}")
def delete_remediation(remediation_id: str, db: Session = Depends(get_db)):
    """Delete a remediation suggestion."""
    try:
        uuid_obj = uuid.UUID(remediation_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID format")
        
    remediation = db.query(models.Remediation).filter(models.Remediation.id == uuid_obj).first()
    if not remediation:
        raise HTTPException(status_code=404, detail="Remediation not found")
        
    db.delete(remediation)
    db.commit()
    return {"status": "success", "message": "Remediation deleted"}

class ZeroDayRequest(BaseModel):
    query: str
    scope: Optional[List[str]] = None  # ["dependencies", "findings", "languages", "all"]

class ZeroDayResponse(BaseModel):
    answer: str
    affected_repositories: List[Dict[str, Any]]
    plan: Optional[Dict[str, Any]] = None
    execution_summary: Optional[List[str]] = None

@router.post("/zero-day", response_model=ZeroDayResponse)
async def analyze_zero_day(
    request: ZeroDayRequest,
    db: Session = Depends(get_db)
):
    """
    Analyze a zero-day vulnerability query with optional scope filtering.
    
    Args:
        request.query: Natural language query about vulnerabilities or technologies
        request.scope: Optional list of data sources to search (dependencies, findings, languages, all)
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    try:
        result = await ai_agent.analyze_zero_day(
            query=request.query,
            db_session=db,
            scope=request.scope
        )
        
        if "error" in result:
             raise HTTPException(status_code=500, detail=result["error"])

        return ZeroDayResponse(
            answer=result.get("answer", "Analysis complete."),
            affected_repositories=result.get("affected_repositories", []),
            plan=result.get("plan"),
            execution_summary=result.get("execution_summary")
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in zero-day analysis: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/zero-day/prompt")
async def get_zero_day_prompt():
    """
    Get the current zero-day analysis planning prompt template.
    """
    # Return the prompt template used in analyze_zero_day
    prompt_template = """You are a Senior Security Analyst Agent analyzing Zero Day vulnerabilities.

User Query: "{query}"
{scope_str}

Available Search Tools:
1. search_dependencies(package_name, version_spec) - Search SBOM/Dependencies for libraries/packages
2. search_findings(query, severity_filter) - Search security findings (CVE, CWE, vulnerabilities)
3. search_languages(language) - Find repositories using specific programming language
4. search_technology(keyword) - Find repos by language or description match

NORMALIZATION RULES (90% Confidence Fuzzy Matching):
- "react.js" or "React" or "ReactJS" → "react"
- "next.js" or "Next.JS" or "NextJS" → "next"
- "log4j" → "log4j-core" or "log4j"
- "python" → "python" or "py"
- Extract CVE IDs (CVE-YYYY-NNNNN format)
- Extract CWE IDs (CWE-NNN format)

SEARCH STRATEGY:
- For package/library vulnerabilities → Use search_dependencies and/or search_findings
- For CVE/CWE queries → Use search_findings
- For technology/language queries → Use search_languages and/or search_technology
- For comprehensive searches → Use multiple tools

Generate a search plan as JSON. Format:
{{
    "thought": "Explain your reasoning and which sources to search",
    "tools": [
        {{"name": "search_dependencies", "args": {{"package_name": "react"}}}},
        {{"name": "search_findings", "args": {{"query": "CVE-2024-12345", "severity_filter": "Critical"}}}},
        {{"name": "search_languages", "args": {{"language": "python"}}}}
    ]
}}

IMPORTANT: Return ONLY the JSON object, no markdown formatting."""
    
    return {"prompt": prompt_template}

class ValidateZDAPromptRequest(BaseModel):
    prompt: str
    test_query: Optional[str] = "Find all repositories using React"

@router.post("/zero-day/prompt/validate")
async def validate_zero_day_prompt(
    request: ValidateZDAPromptRequest,
    db: Session = Depends(get_db)
):
    """
    Validate a zero-day analysis prompt by testing it with a sample query.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
    
    try:
        # Use the test query to validate the prompt
        test_query = request.test_query
        scope_str = "Scope: All sources"
        
        # Replace placeholders in the prompt
        formatted_prompt = request.prompt.replace("{query}", test_query).replace("{scope_str}", scope_str)
        
        # Try to get a plan from the AI
        if hasattr(ai_agent.reasoning_engine.provider, "execute_prompt"):
            result = await ai_agent.reasoning_engine.provider.execute_prompt(formatted_prompt)
            
            # Try to parse as JSON to validate format
            import json
            clean_json = result.strip()
            if clean_json.startswith("```json"):
                clean_json = clean_json.replace("```json", "").replace("```", "")
            elif clean_json.startswith("```"):
                clean_json = clean_json.replace("```", "")
            
            plan = json.loads(clean_json)
            
            return {
                "success": True,
                "response": f"✅ Prompt validated successfully!\n\n**Test Query:** {test_query}\n\n**AI Plan:**\n```json\n{json.dumps(plan, indent=2)}\n```\n\n**Thought:** {plan.get('thought', 'N/A')}\n\n**Tools to Execute:** {len(plan.get('tools', []))}"
            }
        else:
            raise HTTPException(status_code=500, detail="AI Provider does not support prompt execution")
            
    except json.JSONDecodeError as e:
        return {
            "success": False,
            "response": f"❌ **Invalid JSON Response**\n\nThe AI did not return valid JSON. This prompt may need adjustment.\n\n**Error:** {str(e)}\n\n**AI Response:**\n```\n{result[:500]}\n```"
        }
    except Exception as e:
        logger.error(f"Error validating ZDA prompt: {e}")
        return {
            "success": False,
            "response": f"❌ **Validation Error**\n\n{str(e)}"
        }


# Zero Day Analysis Export Endpoints
@router.post("/zero-day/export/pdf")
async def export_zda_pdf(request: dict):
    """Export Zero Day Analysis as PDF."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
        story.append(Paragraph("Zero Day Analysis Report", title_style))
        story.append(Spacer(1, 12))

        # Metadata
        meta_style = styles['Normal']
        story.append(Paragraph(f"<b>Query:</b> {request.get('query', 'N/A')}", meta_style))
        story.append(Paragraph(f"<b>Generated:</b> {request.get('timestamp', 'N/A')}", meta_style))
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        story.append(Paragraph(f"<b>Scope:</b> {scope_str}", meta_style))
        story.append(Spacer(1, 20))

        # Analysis
        story.append(Paragraph("AI Analysis", styles['Heading2']))
        story.append(Spacer(1, 8))
        analysis_text = request.get('analysis', '').replace('\n', '<br/>')
        # Escape XML special characters
        analysis_text = analysis_text.replace('&', '&amp;').replace('<br/>', '<br/>').replace('&amp;', '&')
        try:
            story.append(Paragraph(analysis_text, meta_style))
        except Exception:
            # Fallback to plain text if XML parsing fails
            story.append(Paragraph(request.get('analysis', '').replace('\n', ' ')[:2000], meta_style))
        story.append(Spacer(1, 20))

        # Affected Repositories
        repos = request.get('affected_repositories', [])
        story.append(Paragraph(f"Affected Repositories ({len(repos)})", styles['Heading2']))
        story.append(Spacer(1, 8))

        if repos:
            table_data = [['Repository', 'Reason', 'Source']]
            for repo in repos:
                table_data.append([
                    repo.get('repository', '')[:30],
                    repo.get('reason', 'Context match')[:40],
                    repo.get('source', '-')[:20]
                ])

            table = Table(table_data, colWidths=[2.5*inch, 2.5*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No affected repositories found.", meta_style))

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=zda-analysis.pdf"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation requires reportlab. Install with: pip install reportlab")
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@router.post("/zero-day/export/docx")
async def export_zda_docx(request: dict):
    """Export Zero Day Analysis as DOCX."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('Zero Day Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Query: {request.get('query', 'N/A')}")
        doc.add_paragraph(f"Generated: {request.get('timestamp', 'N/A')}")
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        doc.add_paragraph(f"Scope: {scope_str}")

        # Analysis
        doc.add_heading('AI Analysis', level=1)
        doc.add_paragraph(request.get('analysis', ''))

        # Affected Repositories
        repos = request.get('affected_repositories', [])
        doc.add_heading(f'Affected Repositories ({len(repos)})', level=1)

        if repos:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Repository'
            header_cells[1].text = 'Reason'
            header_cells[2].text = 'Source'

            for repo in repos:
                row_cells = table.add_row().cells
                row_cells[0].text = repo.get('repository', '')
                row_cells[1].text = repo.get('reason', 'Context match')
                row_cells[2].text = repo.get('source', '-')
        else:
            doc.add_paragraph('No affected repositories found.')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=zda-analysis.docx"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX generation requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


# Zero Day Analysis - Repository List Export Endpoints
@router.post("/zero-day/export/repos/pdf")
async def export_zda_repos_pdf(request: dict):
    """Export Zero Day Analysis affected repositories as PDF."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
        story.append(Paragraph("Affected Repositories Report", title_style))
        story.append(Spacer(1, 12))

        # Metadata
        meta_style = styles['Normal']
        story.append(Paragraph(f"<b>Query:</b> {request.get('query', 'N/A')}", meta_style))
        story.append(Paragraph(f"<b>Generated:</b> {request.get('timestamp', 'N/A')}", meta_style))
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        story.append(Paragraph(f"<b>Scope:</b> {scope_str}", meta_style))
        story.append(Paragraph(f"<b>Total Repositories:</b> {request.get('total_repositories', 0)}", meta_style))
        story.append(Spacer(1, 20))

        # Repository Table
        repos = request.get('repositories', [])
        story.append(Paragraph("Repository List", styles['Heading2']))
        story.append(Spacer(1, 8))

        if repos:
            table_data = [['#', 'Repository', 'Reason', 'Source']]
            for idx, repo in enumerate(repos, 1):
                table_data.append([
                    str(idx),
                    repo.get('repository', '')[:25],
                    repo.get('reason', 'Context match')[:35],
                    repo.get('source', '-')[:15]
                ])

            table = Table(table_data, colWidths=[0.4*inch, 2.0*inch, 2.8*inch, 1.3*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('ALIGN', (1, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f5f5')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No affected repositories found.", meta_style))

        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
        story.append(Paragraph("Report generated from Zero Day Analysis", footer_style))

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=affected-repos.pdf"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation requires reportlab. Install with: pip install reportlab")
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@router.post("/zero-day/export/repos/docx")
async def export_zda_repos_docx(request: dict):
    """Export Zero Day Analysis affected repositories as DOCX."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Title
        title = doc.add_heading('Affected Repositories Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Query: {request.get('query', 'N/A')}")
        doc.add_paragraph(f"Generated: {request.get('timestamp', 'N/A')}")
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        doc.add_paragraph(f"Scope: {scope_str}")
        doc.add_paragraph(f"Total Repositories: {request.get('total_repositories', 0)}")

        doc.add_paragraph()  # Spacer

        # Repository List
        doc.add_heading('Repository List', level=1)
        repos = request.get('repositories', [])

        if repos:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header row
            header_cells = table.rows[0].cells
            headers = ['#', 'Repository', 'Reason', 'Source']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Bold header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Data rows
            for idx, repo in enumerate(repos, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = repo.get('repository', '')
                row_cells[2].text = repo.get('reason', 'Context match')
                row_cells[3].text = repo.get('source', '-')
        else:
            doc.add_paragraph('No affected repositories found.')

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('Report generated from Zero Day Analysis')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=affected-repos.docx"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX generation requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


from ..database import get_db
from .. import models
from sqlalchemy.orm import Session
from ..utils.repo_context import get_repo_context, clone_repo_to_temp, cleanup_repo
import uuid

import re

import subprocess
import base64
import os
import tempfile

class ArchitectureRequest(BaseModel):
    project_id: str

class ArchitectureUpdateRequest(BaseModel):
    report: str
    diagram: Optional[str] = None

class PromptRequest(BaseModel):
    project_id: str
    prompt: Optional[str] = None

@router.post("/architecture/prompt")
async def get_architecture_prompt(
    request: ArchitectureRequest,
    db: Session = Depends(get_db),
    settings: settings = Depends(lambda: settings)
):
    """Get the constructed architecture prompt for a project."""
    try:
        p_uuid = uuid.UUID(request.project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == request.project_id).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not project.url:
        raise HTTPException(status_code=400, detail="Project repository URL is missing")

    # Clone repo to temp dir to analyze structure
    try:
        # clone_repo_to_temp creates a new temp dir and returns the path
        temp_dir = clone_repo_to_temp(project.url, settings.GITHUB_TOKEN)
        
        try:
            # Analyze structure
            from ..utils.repo_context import get_repo_structure, get_config_files
            file_structure = get_repo_structure(temp_dir)
            config_files = get_config_files(temp_dir)
            
            # Build prompt
            from ...ai_agent.providers.openai import OpenAIProvider
            provider = OpenAIProvider(api_key=settings.OPENAI_API_KEY, model=settings.AI_MODEL)
            prompt = provider.build_architecture_prompt(project.name, file_structure, config_files, diagrams_index)
            
            return {"prompt": prompt}
        finally:
            cleanup_repo(temp_dir)
            
    except Exception as e:
        logger.error(f"Error building prompt: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/architecture/validate")
async def validate_architecture_prompt(
    request: PromptRequest,
    settings: settings = Depends(lambda: settings)
):
    """Execute a custom architecture prompt."""
    if not request.prompt:
        raise HTTPException(status_code=400, detail="Prompt is required")

    try:
        from ...ai_agent.providers.openai import OpenAIProvider
        provider = OpenAIProvider(api_key=settings.OPENAI_API_KEY, model=settings.AI_MODEL)
        response = await provider.execute_prompt(request.prompt)
        return {"response": response}
        
    except Exception as e:
        logger.error(f"Error validating prompt: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class RefineRequest(BaseModel):
    project_id: str
    code: str

@router.post("/architecture/refine")
async def refine_architecture_diagram(
    request: RefineRequest,
    db: Session = Depends(get_db),
    settings: settings = Depends(lambda: settings)
):
    """Refine diagram code to use correct cloud provider icons."""
    if not request.code:
        raise HTTPException(status_code=400, detail="Code is required")

    try:
        # Fetch the project to get the architecture report
        project = db.query(Project).filter(Project.id == request.project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
            
        report_context = project.architecture_report or ""

        if not ai_agent or not hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
            raise HTTPException(status_code=503, detail="AI provider does not support diagram code fixing")

        # We pass a specific "error" message that acts as the instruction
        instruction = "Refine this code to use the correct cloud provider icons based on the detected technology in the Architecture Report. Enforce the Cloud Provider Preference strictly."

        refined_code = await ai_agent.provider.fix_and_enhance_diagram_code(
            request.code,
            instruction,
            diagrams_index,
            report_context=report_context
        )
        
        # Clean up code block if present
        code_match = re.search(r"```python\n(.*?)```", refined_code, re.DOTALL)
        if code_match:
            refined_code = code_match.group(1).strip()
        else:
            refined_code = refined_code.replace("```python", "").replace("```", "").strip()
            
        return {"code": refined_code}
        
    except Exception as e:
        logger.error(f"Error refining diagram: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class ArchitectureResponse(BaseModel):
    report: str
    diagram: Optional[str] = None # Python code
    image: Optional[str] = None # Base64 PNG

@router.get("/architecture/{project_id}", response_model=ArchitectureResponse)
async def get_architecture(project_id: str, db: Session = Depends(get_db)):
    """Get saved architecture overview for a project."""
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # If we have code but no image (e.g. from old data or failed run), we could try to regen,
    # but for GET let's just return what we have.
    # Actually, we don't store the image in DB, so we need to generate it or store it.
    # Storing base64 in DB might be heavy.
    # Let's regenerate on the fly if missing? No, that's slow.
    # Let's assume we store the code, and maybe we should store the image too?
    # The user asked to "save their edits", implying code edits.
    # If we don't save the image, we have to run the code every time we view? That's slow and risky.
    # Let's add an `architecture_image` column to the DB?
    # Or just return the code and let the frontend trigger a generation if needed?
    # The requirement says "The Python script will then take that specs to create the diagram."
    # Let's stick to generating it on save/generate and returning it.
    # For now, if we don't have the image stored, we might return null and let the UI ask for regen?
    # Or better: let's execute the code if we have it.
    
    image_b64 = None
    if project.architecture_diagram:
        try:
            image_b64 = execute_diagram_code(project.architecture_diagram)
        except Exception as e:
            logger.error(f"Failed to generate image from saved code: {e}")

    return ArchitectureResponse(
        report=project.architecture_report or "",
        diagram=project.architecture_diagram,
        image=image_b64
    )

@router.put("/architecture/{project_id}", response_model=ArchitectureResponse)
async def update_architecture(project_id: str, request: ArchitectureUpdateRequest, db: Session = Depends(get_db)):
    """Update architecture overview (e.g. after user edits)."""
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    project.architecture_report = request.report
    image_b64 = None
    
    if request.diagram is not None:
        project.architecture_diagram = request.diagram
        # Execute code to verify and generate image
        # Execute code to verify and generate image
        try:
            image_b64 = execute_diagram_code(request.diagram)
        except Exception as e:
            logger.warning(f"Updated diagram generation failed: {e}. Attempting auto-fix...")
            try:
                # Auto-fix using the configured AI provider
                if not hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
                    logger.error("AI provider does not support diagram code fixing")
                    raise Exception("AI provider does not support diagram code fixing")

                fixed_code = await ai_agent.provider.fix_and_enhance_diagram_code(request.diagram, str(e), diagrams_index)
                
                # Clean up code block if present
                code_match_fix = re.search(r"```python\n(.*?)```", fixed_code, re.DOTALL)
                if code_match_fix:
                    fixed_code = code_match_fix.group(1).strip()
                else:
                    fixed_code = fixed_code.replace("```python", "").replace("```", "").strip()
                    
                project.architecture_diagram = fixed_code
                image_b64 = execute_diagram_code(fixed_code)
                logger.info("Auto-fix successful")
            except Exception as fix_error:
                logger.error(f"Auto-fix failed: {fix_error}")
                # raise HTTPException(status_code=400, detail=f"Failed to generate diagram: {str(e)}")
        
    db.commit()
    
    return ArchitectureResponse(
        report=project.architecture_report,
        diagram=project.architecture_diagram,
        image=image_b64
    )

def execute_diagram_code(code: str) -> str:
    """Execute Python code to generate diagram and return base64 image."""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Write code to file
        script_path = os.path.join(tmpdir, "diagram_script.py")
        
        # Ensure the code saves to the right filename
        # We look for `filename="architecture_diagram"` or similar and enforce output path
        # Or we just run it and look for the .png file.
        # The prompt instructs `filename="architecture_diagram"`.
        # So we expect `architecture_diagram.png` in the CWD.
        
        with open(script_path, "w") as f:
            f.write(code)
            
        # Run script
        try:
            subprocess.check_output(
                ["python3", script_path],
                cwd=tmpdir,
                stderr=subprocess.STDOUT,
                timeout=30
            )
        except subprocess.CalledProcessError as e:
            raise Exception(f"Script execution failed: {e.output.decode()}")
            
        # Find PNG
        png_path = os.path.join(tmpdir, "architecture_diagram.png")
        if not os.path.exists(png_path):
             # Try finding any png
            files = [f for f in os.listdir(tmpdir) if f.endswith('.png')]
            if files:
                png_path = os.path.join(tmpdir, files[0])
            else:
                raise Exception("No PNG image generated by the script")
            
        with open(png_path, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")

@router.post("/architecture", response_model=ArchitectureResponse)
async def generate_architecture(request: ArchitectureRequest, db: Session = Depends(get_db)):
    """Generate an architecture overview for a project."""
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
        
    # Get project
    try:
        p_uuid = uuid.UUID(request.project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == request.project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if not project.url:
        raise HTTPException(status_code=400, detail="Project repository URL is missing")
        
    repo_path = None
    try:
        # Clone repo to temp
        # Use GITHUB_TOKEN from settings if available
        token = settings.GITHUB_TOKEN
        repo_path = clone_repo_to_temp(project.url, token)
        
        # Get context
        structure, configs = get_repo_context(repo_path)
        
        # Generate overview
        full_response = await ai_agent.generate_architecture_overview(
            repo_name=project.name,
            file_structure=structure,
            config_files=configs
        )
        
        # Parse Python code from response
        diagram_code = None
        report = full_response
        image_b64 = None
        
        code_match = re.search(r"```python\n(.*?)```", full_response, re.DOTALL)
        if code_match:
            diagram_code = code_match.group(1).strip()
            # Remove the code block from the report
            report = full_response.replace(code_match.group(0), "").strip()
            
            # Generate image
            # Generate image
            try:
                image_b64 = execute_diagram_code(diagram_code)
            except Exception as e:
                logger.warning(f"Initial diagram generation failed: {e}. Attempting auto-fix...")
                try:
                    # Auto-fix using the configured AI provider
                    if not hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
                        logger.error("AI provider does not support diagram code fixing")
                        raise Exception("AI provider does not support diagram code fixing")

                    fixed_code = await ai_agent.provider.fix_and_enhance_diagram_code(diagram_code, str(e), diagrams_index)
                    
                    # Clean up code block if present
                    code_match_fix = re.search(r"```python\n(.*?)```", fixed_code, re.DOTALL)
                    if code_match_fix:
                        fixed_code = code_match_fix.group(1).strip()
                    else:
                        fixed_code = fixed_code.replace("```python", "").replace("```", "").strip()
                        
                    diagram_code = fixed_code
                    image_b64 = execute_diagram_code(diagram_code)
                    logger.info("Auto-fix successful")
                except Exception as fix_error:
                    logger.error(f"Auto-fix failed: {fix_error}")
                    # We still save the code so user can fix it
            
        # Save to DB
        project.architecture_report = report
        project.architecture_diagram = diagram_code
        db.commit()
        
        return ArchitectureResponse(
            report=report,
            diagram=diagram_code,
            image=image_b64
        )
        
    except Exception as e:
        logger.error(f"Error generating architecture: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if repo_path:
            cleanup_repo(repo_path)

class VersionCreateRequest(BaseModel):
    description: Optional[str] = None

class VersionResponse(BaseModel):
    id: str
    version_number: int
    created_at: str
    description: Optional[str] = None
    
class VersionDetailResponse(VersionResponse):
    report: str
    diagram: Optional[str] = None

@router.post("/architecture/{project_id}/versions", response_model=VersionResponse)
async def create_architecture_version(
    project_id: str, 
    request: VersionCreateRequest, 
    db: Session = Depends(get_db)
):
    """Save current architecture state as a new version."""
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if not project.architecture_report:
        raise HTTPException(status_code=400, detail="No architecture report to save")
        
    # Get next version number
    last_version = db.query(func.max(models.ArchitectureVersion.version_number))\
        .filter(models.ArchitectureVersion.repository_id == project.id)\
        .scalar() or 0
        
    new_version = models.ArchitectureVersion(
        repository_id=project.id,
        version_number=last_version + 1,
        report_content=project.architecture_report,
        diagram_code=project.architecture_diagram,
        description=request.description
    )
    
    db.add(new_version)
    db.commit()
    db.refresh(new_version)
    
    return VersionResponse(
        id=str(new_version.id),
        version_number=new_version.version_number,
        created_at=new_version.created_at.isoformat(),
        description=new_version.description
    )

@router.get("/architecture/{project_id}/versions", response_model=List[VersionResponse])
async def list_architecture_versions(project_id: str, db: Session = Depends(get_db)):
    """List all architecture versions for a project."""
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    versions = db.query(models.ArchitectureVersion)\
        .filter(models.ArchitectureVersion.repository_id == project.id)\
        .order_by(models.ArchitectureVersion.version_number.desc())\
        .all()
        
    return [
        VersionResponse(
            id=str(v.id),
            version_number=v.version_number,
            created_at=v.created_at.isoformat(),
            description=v.description
        ) for v in versions
    ]

@router.get("/architecture/{project_id}/versions/{version_id}", response_model=VersionDetailResponse)
async def get_architecture_version(project_id: str, version_id: str, db: Session = Depends(get_db)):
    """Get details of a specific architecture version."""
    try:
        v_uuid = uuid.UUID(version_id)
        version = db.query(models.ArchitectureVersion).filter(models.ArchitectureVersion.id == v_uuid).first()
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid version ID")
        
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
        
    return VersionDetailResponse(
        id=str(version.id),
        version_number=version.version_number,
        created_at=version.created_at.isoformat(),
        description=version.description,
        report=version.report_content,
        diagram=version.diagram_code
    )

@router.post("/architecture/{project_id}/restore/{version_id}")
async def restore_architecture_version(project_id: str, version_id: str, db: Session = Depends(get_db)):
    """Restore a previous architecture version."""
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    try:
        v_uuid = uuid.UUID(version_id)
        version = db.query(models.ArchitectureVersion).filter(models.ArchitectureVersion.id == v_uuid).first()
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid version ID")
        
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
        
    # Overwrite current state
    project.architecture_report = version.report_content
    project.architecture_diagram = version.diagram_code
    
    db.commit()
    
    return {"status": "success", "message": f"Restored version {version.version_number}"}
